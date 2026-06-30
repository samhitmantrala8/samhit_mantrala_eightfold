from __future__ import annotations

from pathlib import Path

from docx import Document

from backend.transformer.extractors.notes_extractor import extract_notes_from_text
from backend.transformer.facts import ExtractionBundle


def extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def extract_docx(path: Path, use_llm: bool = False) -> ExtractionBundle:
    try:
        text = extract_docx_text(path)
    except Exception as exc:
        return ExtractionBundle([], [f"{path.name}: failed to parse DOCX: {exc}"])
    if not text.strip():
        return ExtractionBundle([], [f"{path.name}: no extractable DOCX text found"])
    return extract_notes_from_text(text, f"docx:{path.name}", use_llm=use_llm)
